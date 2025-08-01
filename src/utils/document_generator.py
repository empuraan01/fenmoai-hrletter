"""
Document generation utilities for PDF and DOCX formats
"""

import io
import re
from datetime import datetime
from typing import Dict, Any

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentGenerator:
    """Generate professional documents in PDF and DOCX formats"""
    
    def __init__(self):
        self.company_name = "Company ABC"
        self.company_address = "123 Business Park, Tech City, IN 560001"
        
    def generate_pdf(self, offer_letter_text: str, employee_name: str) -> bytes:
        """Generate PDF version of the offer letter"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab library not available. Install with: pip install reportlab>=4.0.0")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        story = []
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            alignment=TA_LEFT
        )
        
        story.append(Paragraph(f"<b>{self.company_name}</b>", title_style))
        story.append(Paragraph(self.company_address, body_style))
        story.append(Spacer(1, 20))
        
        current_date = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(f"<b>Date:</b> {current_date}", body_style))
        story.append(Spacer(1, 12))
        
        lines = offer_letter_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
                continue
                
            if any(char in line for char in '📄🎯💰📋🏢📝✍️') or line.isupper():
                if len(line) > 50:
                    story.append(Paragraph(line, body_style))
                else:
                    story.append(Paragraph(f"<b>{line}</b>", header_style))
            else:
                story.append(Paragraph(line, body_style))
        
        story.append(Spacer(1, 30))
        story.append(Paragraph("<b>Sincerely,</b>", body_style))
        story.append(Spacer(1, 40))
        
        story.append(Paragraph("_______________________", body_style))
        story.append(Paragraph("HR Manager", body_style))
        story.append(Paragraph(f"{self.company_name}", body_style))
        
        doc.build(story)
        
        pdf_data = buffer.getvalue()
        buffer.close()
        return pdf_data
    
    def generate_docx(self, offer_letter_text: str, employee_name: str) -> bytes:
        """Generate DOCX version of the offer letter"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available. Install with: pip install python-docx>=1.1.0")
        
        doc = Document()
        
        header = doc.add_heading(self.company_name, 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        address_para = doc.add_paragraph(self.company_address)
        address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        current_date = datetime.now().strftime("%B %d, %Y")
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Date: {current_date}")
        date_run.bold = True
        
        doc.add_paragraph()

        lines = offer_letter_text.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
                
            clean_line = re.sub(r'[^\w\s\-.,;:()[]{}$₹%/]', '', line)
            
            if line.isupper() or any(keyword in line.lower() for keyword in ['appointment details', 'compensation structure', 'terms and conditions']):
                if len(clean_line) > 50:
                    doc.add_paragraph(clean_line)
                else:
                    doc.add_heading(clean_line, level=2)
            else:
                doc.add_paragraph(clean_line)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        sig_para = doc.add_paragraph()
        sig_run = sig_para.add_run("Sincerely,")
        sig_run.bold = True
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("_______________________")
        doc.add_paragraph("HR Manager")
        doc.add_paragraph(f"{self.company_name}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        docx_data = buffer.getvalue()
        buffer.close()
        return docx_data
    
    def get_available_formats(self) -> Dict[str, bool]:
        """Get available document formats"""
        return {
            'txt': True,
            'pdf': REPORTLAB_AVAILABLE,
            'docx': DOCX_AVAILABLE
        } 